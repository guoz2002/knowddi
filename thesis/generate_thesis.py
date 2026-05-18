"""
generate_thesis.py
主论文生成脚本：读取各章节 Python 模块中的内容，按山东科技大学毕业论文格式生成 Word 文档
运行方式：python generate_thesis.py
"""

import os
import sys
import importlib.util

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ──────────────────────────────────────────────
# 路径配置
# ──────────────────────────────────────────────
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
CHAPTERS_DIR = os.path.join(SCRIPT_DIR, "chapters")
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "output", "基于知识图谱增强的药物相互作用预测研究.docx")
os.makedirs(os.path.join(SCRIPT_DIR, "output"), exist_ok=True)

# ──────────────────────────────────────────────
# 动态加载章节模块
# ──────────────────────────────────────────────
def load_chapter_module(filename):
    """从 chapters/ 目录中按文件名动态加载模块"""
    filepath = os.path.join(CHAPTERS_DIR, filename)
    spec = importlib.util.spec_from_file_location(filename[:-3], filepath)
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod


CHAPTER_FILES = [
    "ch01_introduction.py",
    "ch02_background.py",
    # ch03_baseline.py 已废弃：原第三章内容已分散并入第四、第五章 -- 修改新增
    "ch04_ablation.py",
    "ch05_lightweight.py",
    "ch06_conclusion.py",
    "references_acknowledgement.py",
]

# ──────────────────────────────────────────────
# 字体与段落格式辅助函数（依据山东科技大学本科毕业论文规范）
# ──────────────────────────────────────────────

def set_font(run, name_cn="宋体", name_en="Times New Roman", size_pt=12, bold=False):
    """统一设置中英文字体"""
    run.font.name = name_en
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    # 设置中文字体
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), name_cn)
    rFonts.set(qn('w:ascii'), name_en)
    rFonts.set(qn('w:hAnsi'), name_en)
    existing = rPr.find(qn('w:rFonts'))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)


def set_paragraph_format(para,
                          alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                          first_line_indent_cm=0.85,
                          space_before_pt=0,
                          space_after_pt=0,
                          line_spacing_pt=22):
    """统一设置段落格式（首行缩进、行距等）"""
    fmt = para.paragraph_format
    fmt.alignment = alignment
    if first_line_indent_cm:
        fmt.first_line_indent = Cm(first_line_indent_cm)
    fmt.space_before = Pt(space_before_pt)
    fmt.space_after = Pt(space_after_pt)
    from docx.shared import Pt as DocxPt
    from docx.oxml.ns import qn
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:line'), str(int(line_spacing_pt * 20)))
    spacing.set(qn('w:lineRule'), 'exact')
    existing = pPr.find(qn('w:spacing'))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(spacing)


import re

RED = RGBColor(0xFF, 0x00, 0x00)
GREEN = RGBColor(0x00, 0x8B, 0x00)  # 新增内容标注色（深绿）
BLUE = RGBColor(0x1F, 0x49, 0xD8)  # 二轮新增技术细节标注色（蓝色） -- 修改新增
YELLOW = RGBColor(0xCC, 0x99, 0x00)  # 三轮新增图文衔接句标注色（深黄/豁黄，保证打印可读） -- 修改新增
GRAY = RGBColor(0x80, 0x80, 0x80)  # 四轮补充内容标注色（灰色） -- 修改新增

# 匹配引用标记 [数字] 或 [数字, 数字] 等
REF_PATTERN = re.compile(r'\[\d+(?:[,，]\d+)*\]')

# 需要插入图/代码的占位关键词（每个关键词全文只插一次）
# 格式：(唯一触发词, 占位说明)
FIGURE_KEYWORDS = [
    ("总体数据流", "【待插图：KnowDDI整体架构与数据流示意图】"),
    ("有向子图提取（DIG）模块", "【待插图：DIG子图抽取算法示意图】"),
    ("EdgeGateNetwork", "【待插图：EdgeGateNetwork结构或代码片段截图】"),
    ("DrugBank消融实验结果\n\n表4-1", "【待插图：DrugBank消融实验各指标对比柱状图】"),
    ("TWOSIDES消融实验结果\n\n表4-2", "【待插图：TWOSIDES消融实验各指标对比柱状图】"),
    ("不同子图稀疏强度下的性能变化（DrugBank）\n\n表5-1", "【待插图：DrugBank子图稀疏强度-精度/效率曲线图】"),
    ("不同特征压缩强度下的性能变化（DrugBank）\n\n表5-2", "【待插图：DrugBank特征压缩强度-精度/效率曲线图】"),
    ("DrugBank轻量化方案综合对比（相对于基线的变化率", "【待插图：DrugBank轻量化方案综合对比图（精度损失 vs 资源节省）】"),
]
# 记录已触发过的关键词，避免重复插入
_triggered_keywords: set = set()


def add_body_paragraph_with_highlights(doc, text, indent=True, use_green=False, use_blue=False, use_yellow=False, use_gray=False):
    """
    添加正文段落：
    - 引用标记 [N] 标红
    - use_green=True 时正文内容显示为绿色（一轮新增）
    - use_blue=True  时正文内容显示为蓝色（二轮新增技术细节） -- 修改新增
    - use_yellow=True时正文内容显示为黄色（三轮新增图文衔接句） -- 修改新增
    - use_gray=True  时正文内容显示为灰色（四轮补充内容） -- 修改新增
    - 文本内出现 {{B:...}} 行内片段时，该片段单独以蓝色渲染（行内二轮新增） -- 修改新增
    - 文本内出现 {{Y:...}} 行内片段时，该片段单独以黄色渲染（行内三轮新增） -- 修改新增
    - 文本内出现 {{G:...}} 行内片段时，该片段单独以灰色渲染（行内四轮补充） -- 修改新增
    - 遇到图片占位关键词后追加红色占位提示段
    """
    para = doc.add_paragraph()
    set_paragraph_format(para, first_line_indent_cm=0.85 if indent else 0)

    # 拆分文本：普通片段 与 引用标记交替
    parts = REF_PATTERN.split(text.strip())
    refs = REF_PATTERN.findall(text.strip())

    # 行内色片段的正则（蓝色 + 黄色 + 灰色） -- 修改新增
    inline_color_pat = re.compile(r"\{\{([BYG]):(.*?)\}\}", flags=re.DOTALL)

    def _emit_segment(seg_text, force_blue=False):
        """按行内色片段切分并写入 run -- 修改新增"""
        last = 0
        for m in inline_color_pat.finditer(seg_text):
            pre = seg_text[last:m.start()]
            if pre:
                r = para.add_run(pre)
                set_font(r, name_cn="宋体", name_en="Times New Roman", size_pt=12)
                if force_blue or use_blue:
                    r.font.color.rgb = BLUE
                elif use_yellow:
                    r.font.color.rgb = YELLOW
                elif use_gray:
                    r.font.color.rgb = GRAY
                elif use_green:
                    r.font.color.rgb = GREEN
            tag, inner_text = m.group(1), m.group(2)
            if inner_text:
                r = para.add_run(inner_text)
                set_font(r, name_cn="宋体", name_en="Times New Roman", size_pt=12)
                if tag == 'Y':
                    r.font.color.rgb = YELLOW
                elif tag == 'G':
                    r.font.color.rgb = GRAY
                else:
                    r.font.color.rgb = BLUE
            last = m.end()
        tail = seg_text[last:]
        if tail:
            r = para.add_run(tail)
            set_font(r, name_cn="宋体", name_en="Times New Roman", size_pt=12)
            if force_blue or use_blue:
                r.font.color.rgb = BLUE
            elif use_yellow:
                r.font.color.rgb = YELLOW
            elif use_gray:
                r.font.color.rgb = GRAY
            elif use_green:
                r.font.color.rgb = GREEN

    for i, part in enumerate(parts):
        if part:
            _emit_segment(part)
        if i < len(refs):
            ref_run = para.add_run(refs[i])
            set_font(ref_run, name_cn="宋体", name_en="Times New Roman", size_pt=12)
            ref_run.font.color.rgb = RED

    # 检查是否需要在本段后插入红色占位提示（每个关键词全文只插一次）
    for keyword, placeholder in FIGURE_KEYWORDS:
        # 使用简单关键词（取第一行）做匹配
        simple_key = keyword.split("\n")[0]
        if simple_key in text and simple_key not in _triggered_keywords:
            _triggered_keywords.add(simple_key)
            ph_para = doc.add_paragraph()
            ph_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ph_para.paragraph_format.first_line_indent = Pt(0)
            ph_run = ph_para.add_run(placeholder)
            set_font(ph_run, name_cn="宋体", name_en="Times New Roman", size_pt=12)
            ph_run.font.color.rgb = RED
            ph_run.font.bold = True
            break

    return para



def add_chapter_heading(doc, text):
    """添加章标题：黑体三号（16pt），居中，段前段后各12pt"""
    para = doc.add_paragraph()
    fmt = para.paragraph_format
    fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt.space_before = Pt(12)
    fmt.space_after = Pt(12)
    fmt.first_line_indent = Pt(0)
    run = para.add_run(text)
    set_font(run, name_cn="黑体", name_en="Arial", size_pt=16, bold=True)
    return para


def add_section_heading(doc, text, level=2):
    """添加节标题（一级：小三14pt黑体；二级：四号14pt黑体；三级：小四12pt黑体）"""
    pt = {2: 14, 3: 14, 4: 12}.get(level, 12)
    para = doc.add_paragraph()
    fmt = para.paragraph_format
    fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt.space_before = Pt(6)
    fmt.space_after = Pt(6)
    fmt.first_line_indent = Pt(0)
    run = para.add_run(text)
    set_font(run, name_cn="黑体", name_en="Arial", size_pt=pt, bold=True)
    return para


def add_page_break(doc):
    doc.add_page_break()


# ──────────────────────────────────────────────
# 表格辅助
# ──────────────────────────────────────────────

def parse_markdown_table(text):
    """解析 Markdown 表格，返回 (headers, rows) 二元组"""
    lines = [l.strip() for l in text.strip().splitlines() if l.strip()]
    table_lines = [l for l in lines if l.startswith("|")]
    if len(table_lines) < 2:
        return None, None
    headers = [h.strip() for h in table_lines[0].strip("|").split("|")]
    rows = []
    for row_line in table_lines[2:]:  # 跳过分隔行
        cells = [c.strip() for c in row_line.strip("|").split("|")]
        rows.append(cells)
    return headers, rows


def add_docx_table(doc, headers, rows, caption=None):
    """在 Word 中添加规范格式表格"""
    if caption:
        cap_para = doc.add_paragraph()
        cap_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_para.paragraph_format.first_line_indent = Pt(0)
        cap_run = cap_para.add_run(caption)
        set_font(cap_run, name_cn="宋体", name_en="Times New Roman", size_pt=10.5, bold=True)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # 表头
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, name_cn="宋体", name_en="Times New Roman", size_pt=10.5, bold=True)

    # 数据行
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(cell_text.replace("**", "").replace("*", ""))
            set_font(run, name_cn="宋体", name_en="Times New Roman", size_pt=10.5)

    # 表后空行
    doc.add_paragraph()


# ──────────────────────────────────────────────
# 内容解析与渲染
# ──────────────────────────────────────────────

def render_section_content(doc, content_text, use_green=False):
    """
    将节正文（可能包含 Markdown 表格、小节标题、正文段落）渲染到 Word
    """
    lines = content_text.split("\n")
    i = 0
    table_buffer = []
    in_table = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 检测 Markdown 表格开始
        if stripped.startswith("|"):
            in_table = True
            table_buffer.append(stripped)
            i += 1
            continue

        # 表格结束
        if in_table and not stripped.startswith("|"):
            in_table = False
            if table_buffer:
                headers, rows = parse_markdown_table("\n".join(table_buffer))
                if headers and rows:
                    add_docx_table(doc, headers, rows)
                table_buffer = []

        # 三级节标题（如 "2.1.1 ..."）
        if stripped and stripped[0].isdigit() and "." in stripped[:6] and stripped[3:4] in [".", " "]:
            # 粗略判断是否是子节标题
            parts = stripped.split(" ", 1)
            if len(parts) == 2 and parts[0].replace(".", "").isdigit():
                num_dots = parts[0].count(".")
                level = min(num_dots + 1, 4)
                add_section_heading(doc, stripped, level=level)
                i += 1
                continue

        # 空行：跳过
        if not stripped:
            i += 1
            continue

        # 普通正文
        # 支持段落级绿色标注：以 [GREEN] 开头的段落单独以绿色渲染（修改新增）
        para_use_green = use_green
        para_use_blue = False  # 二轮新增技术细节标注 -- 修改新增
        para_use_yellow = False  # 三轮新增图文衔接句标注 -- 修改新增
        para_text = stripped
        if para_text.startswith("[GREEN]"):
            para_use_green = True
            para_text = para_text[len("[GREEN]"):].lstrip()
        # 支持二轮新增技术细节段落：以 [BLUE] 开头的段落以蓝色渲染 -- 修改新增
        if para_text.startswith("[BLUE]"):
            para_use_blue = True
            para_use_green = False
            para_text = para_text[len("[BLUE]"):].lstrip()
        # 支持三轮新增图文衔接句段落：以 [YELLOW] 开头的段落以黄色渲染 -- 修改新增
        if para_text.startswith("[YELLOW]"):
            para_use_yellow = True
            para_use_green = False
            para_use_blue = False
            para_text = para_text[len("[YELLOW]"):].lstrip()
        # 支持四轮补充内容段落：以 [GRAY] 开头的段落以灰色渲染 -- 修改新增
        para_use_gray = False
        if para_text.startswith("[GRAY]"):
            para_use_gray = True
            para_use_green = False
            para_use_blue = False
            para_use_yellow = False
            para_text = para_text[len("[GRAY]"):].lstrip()
        # 支持代码截图占位：以 [CODE] 开头的段落渲染为居中红色粗体提示段（修改新增）
        if para_text.startswith("[CODE]"):
            ph_para = doc.add_paragraph()
            ph_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ph_para.paragraph_format.first_line_indent = Pt(0)
            ph_run = ph_para.add_run(para_text[len("[CODE]"):].lstrip())
            set_font(ph_run, name_cn="宋体", name_en="Times New Roman", size_pt=12, bold=True)
            ph_run.font.color.rgb = RED
            i += 1
            continue
        # 支持二轮新增图/代码占位：以 [BLUEFIG] 开头的段落渲染为居中蓝色粗体占位段 -- 修改新增
        if para_text.startswith("[BLUEFIG]"):
            ph_para = doc.add_paragraph()
            ph_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ph_para.paragraph_format.first_line_indent = Pt(0)
            ph_run = ph_para.add_run(para_text[len("[BLUEFIG]"):].lstrip())
            set_font(ph_run, name_cn="宋体", name_en="Times New Roman", size_pt=12, bold=True)
            ph_run.font.color.rgb = BLUE
            i += 1
            continue
        add_body_paragraph_with_highlights(doc, para_text, use_green=para_use_green, use_blue=para_use_blue, use_yellow=para_use_yellow, use_gray=para_use_gray)
        i += 1

    # 处理文件末尾的悬空表格
    if table_buffer:
        headers, rows = parse_markdown_table("\n".join(table_buffer))
        if headers and rows:
            add_docx_table(doc, headers, rows)


# ──────────────────────────────────────────────
# 页面设置
# ──────────────────────────────────────────────

def setup_page(doc):
    """设置 A4 页面，页边距 上3cm 下2.5cm 左3cm 右2.5cm"""
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2.5)


# ──────────────────────────────────────────────
# 封面
# ──────────────────────────────────────────────

def add_cover(doc):
    """生成封面页"""
    for _ in range(4):
        doc.add_paragraph()

    # 校名
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run("山东科技大学")
    set_font(r, "黑体", "Arial", 22, bold=True)

    doc.add_paragraph()

    # 大标题
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run("本科毕业设计（论文）")
    set_font(r, "黑体", "Arial", 22, bold=True)

    for _ in range(3):
        doc.add_paragraph()

    # 论文题目
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run("基于知识图谱增强的药物相互作用预测研究")
    set_font(r, "黑体", "Arial", 18, bold=True)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run("——图结构学习机制解耦与轻量化重构")
    set_font(r, "宋体", "Times New Roman", 14)

    for _ in range(4):
        doc.add_paragraph()

    # 作者信息表格
    info_items = [
        ("学    院", "计算机科学与工程学院"),
        ("专    业", "计算机科学与技术"),
        ("班    级", "2021级×班"),
        ("学生姓名", "××××"),
        ("学    号", "202×××××××"),
        ("指导教师", "××× 副教授"),
        ("完成日期", "2025年6月"),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        r = p.add_run(f"{label}：{value}")
        set_font(r, "宋体", "Times New Roman", 14)

    add_page_break(doc)


# ──────────────────────────────────────────────
# 摘要
# ──────────────────────────────────────────────

def add_abstract_cn(doc):
    """中文摘要页"""
    add_chapter_heading(doc, "摘  要")

    abstract_text = (
        "药物-药物相互作用（Drug-Drug Interaction, DDI）的准确预测对临床联合用药安全和新药研发具有重要意义。"
        "现有基于图神经网络（GNN）的DDI预测方法虽已引入外部生物医学知识图谱（KG）来弥补已知DDI样本的稀缺性，"
        "但图结构学习中的\u201c去噪\u201d与\u201c补全\u201d两类操作对预测性能的独立贡献尚未得到系统性定量分析；"
        "与此同时，现有模型参数量大、推理延时高，难以满足医疗边缘端部署的需求。"
        "{{Y:以KnowDDI为例，其图结构学习模块通过EdgeGateNetwork同时学习gate/denoise/completion三路打分并融合为边权，"
        "背景知识图谱采用Hetionet v1.0来增强稀疏DDI数据中的实体关联，"
        "去噪分支仅对子图中已有的边重加权、补全分支仅为缺失边补分，"
        "二者共享隐藏层权重却始终被联合训练，各自的独立贡献从未被单独量化；"
        "模型在两个基准数据集上的推理显存峰值达911~1667 MiB，难以直接满足医疗边缘场景的部署要求。}}"
    )
    add_body_paragraph_with_highlights(doc, abstract_text)

    text2 = (
        "本文以KnowDDI模型为基础实现平台，围绕以下两条主线开展研究："
        "（1）通过设计严格的受控消融实验，构建baseline、仅去噪、仅补全和去噪+补全四组变体，"
        "在DrugBank（多分类）与TWOSIDES（多标签）两个基准数据集上对图结构学习中的去噪与补全机制进行定量解耦，"
        "揭示两类操作的差异化边际贡献与数据集相关的适用边界；"
        "（2）提出\u201c拓扑-参数\u201d协同压缩策略，通过子图稀疏采样与多维特征压缩的联合优化，"
        "系统评估不同轻量化方案在精度、推理时延与显存占用之间的折中关系。"
        "{{Y:消融研究中，四组变体通过新增命令行参数--gsl_mode统一切换去噪/补全分支的激活状态，"
        "参数量完全一致，消融对比不受模型容量差异干扰；"
        "实验在DrugBank（1710种药物、86类DDI关系）与TWOSIDES（645种药物、200类DDI关系）两个数据集上，"
        "各以3个固定随机种子重复运行，采用与原论文一致的实验协议，"
        "基线复现结果与原报告值的偏差均控制在合理范围内。"
        "轻量化研究中，拓扑压缩入口为子图采样的max_nodes_per_hop与max_links（60%保留比例），"
        "参数压缩入口为emb_dim/gsl_rel_emb_dim/MLP_hidden_dim三维同步缩减（50%保留比例），"
        "两类入口正交独立，方案间比较不受参数配置歧义影响。}}"
    )
    add_body_paragraph_with_highlights(doc, text2)

    text3 = (
        "实验结果表明：在DrugBank数据集上，去噪机制更有利于提升类别均衡意义下的预测性能（Macro-F1），"
        "补全机制更有利于增强整体预测一致性（ACC、Cohen's κ）；在TWOSIDES数据集上，"
        "原始基线模型表现最优，补全模块激活比例极低（0.62%~0.79%），说明两类机制的适用边界具有显著的数据集依赖性。"
        "{{Y:具体而言，DrugBank上Denoise-only变体取得四组最高Macro-F1（91.52%），"
        "Completion-only变体的ACC（92.89%）与Cohen's κ（91.57%）略高于基线，"
        "Denoise+Completion联合方案在ACC（92.99%）和κ（91.68%）上达到全局最优，"
        "但其Macro-F1（90.78%）受补全引入的长尾类别误判边拖累，低于单独去噪变体；"
        "TWOSIDES上基线AUROC为95.44%、AUPRC为94.11%，三个引入GSL操作的变体均低于基线，"
        "补全激活比例随训练逐步从约1.1%自发收敛至0.6%~0.8%，"
        "模型在该数据集上自发地将补全边比例压到极低水平，无需外部干预。}}"
        "在轻量化实验中，仅特征压缩方案（50%压缩比）在DrugBank上以仅0.58%的Macro-F1损失实现了约48%的显存节省，"
        "是精度损失最小的单维度轻量化策略；协同优化方案相比单一特征压缩的额外资源收益有限，"
        "但为极限资源压缩场景提供了更大的压缩空间。"
        "{{Y:仅子图稀疏方案在DrugBank上Macro-F1损失1.73%、显存节省43.9%，"
        "两种方案的推理时长节省均接近48%，但协同压缩在DrugBank上未获得相比特征压缩更明显的时延收益，"
        "归因分析表明GraphSAGE节点嵌入索引步骤的耗时与emb_dim呈线性相关而与节点数仅呈次线性相关，"
        "emb_dim压缩已率先消除该瓶颈；TWOSIDES上三类轻量化方案的效率收益均不超过10%，"
        "说明该数据集的计算开销主要受样本规模驱动，单纯的拓扑或特征压缩难以带来显著的推理提速。}}"
    )
    add_body_paragraph_with_highlights(doc, text3)

    # 新增：可解释性段落（绿色标注）
    text4_prefix = ""
    text4_new = (
        "此外，本文选取Reboxetine-Atomoxetine与Atropine-Scopolamine两个典型药物对，"
        "通过分析基线模型在知识子图中保留的高权重边与推理路径，验证了KnowDDI模型基于知识子图学习的可解释预测能力。"
        "前者呈现出经由桥接化合物DB01146的清晰双跳推理路径，后者则体现了模型对局部高权重邻域的综合聚合推理。"
        "上述案例表明，知识子图不仅是提升预测精度的结构工具，也是模型预测结果的生物医学解释载体。"
    )
    p_new = doc.add_paragraph()
    p_new.paragraph_format.first_line_indent = Pt(24)
    p_new.paragraph_format.space_before = Pt(0)
    p_new.paragraph_format.space_after = Pt(0)
    r_new = p_new.add_run(text4_new)
    set_font(r_new, "宋体", "Times New Roman", 12)
    r_new.font.color.rgb = GREEN  # 绿色标注新增内容

    doc.add_paragraph()
    kw_para = doc.add_paragraph()
    kw_para.paragraph_format.first_line_indent = Pt(0)
    r = kw_para.add_run("关键词：")
    set_font(r, "黑体", "Arial", 12, bold=True)
    r2 = kw_para.add_run("药物-药物相互作用预测；图神经网络；知识图谱；图结构学习；轻量化模型；消融实验；可解释性")
    set_font(r2, "宋体", "Times New Roman", 12)

    add_page_break(doc)


def add_abstract_en(doc):
    """英文摘要页"""
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("ABSTRACT")
    set_font(r, "宋体", "Arial", 16, bold=True)

    texts_base = [
        ("Accurate prediction of Drug-Drug Interactions (DDI) is of critical importance for "
         "clinical polypharmacy safety and drug discovery. Although existing GNN-based DDI prediction "
         "methods have incorporated external biomedical knowledge graphs (KG) to compensate for the "
         "scarcity of labeled DDI samples, the independent contributions of graph structure learning "
         "operations—namely denoising and completion—to predictive performance have not been "
         "systematically quantified. Meanwhile, existing models suffer from large parameter counts "
         "and high inference latency, making them difficult to deploy on resource-constrained "
         "medical edge devices."),
        ("This thesis uses the KnowDDI model as the implementation platform and conducts research "
         "along two main lines: (1) Through rigorous controlled ablation experiments, four model "
         "variants—baseline, denoise-only, completion-only, and denoise+completion—are constructed "
         "to quantitatively decouple the denoising and completion mechanisms on DrugBank "
         "(multi-class) and TWOSIDES (multi-label) benchmark datasets, revealing their "
         "differentiated marginal contributions and dataset-dependent applicability boundaries; "
         "(2) A 'topology-parameter' co-compression strategy is proposed, and the trade-off "
         "among accuracy, inference latency, and GPU memory usage of various lightweight schemes "
         "is systematically evaluated."),
        ("Experimental results show that on DrugBank, the denoising mechanism more effectively "
         "improves category-balanced prediction performance (Macro-F1), while the completion "
         "mechanism more effectively enhances overall prediction consistency (ACC, Cohen's κ). "
         "On TWOSIDES, the original baseline achieves the best performance and the completion "
         "module activation ratio is extremely low (0.62%–0.79%), indicating that the "
         "applicability of the two mechanisms is significantly dataset-dependent. In lightweight "
         "experiments, the feature-only compression scheme (50% ratio) achieves approximately "
         "48% memory savings on DrugBank with only 0.58% Macro-F1 loss, making it the "
         "best-trade-off single-dimension lightweight strategy."),
    ]

    # 各段扩充内容（黄色）-- 修改新增
    texts_ext = [
        ("{{Y:Specifically, KnowDDI employs an EdgeGateNetwork to jointly learn gate/denoise/completion "
         "three-way branch scores over candidate subgraph edges, using Hetionet v1.0 as the background "
         "knowledge graph. The denoising branch re-weights existing edges while the completion branch "
         "scores absent ones; though they share a hidden layer, their individual contributions have "
         "never been isolated. Peak GPU memory usage reaches 911–1667 MiB across the two benchmarks, "
         "falling short of typical medical edge deployment constraints.}}"),
        ("{{Y:In the ablation study, all four variants are controlled via a single command-line argument "
         "--gsl_mode that switches the denoise/completion branch activation while keeping parameter counts "
         "identical, eliminating confounding from model capacity differences. Experiments are run on "
         "DrugBank (1710 drugs, 86 DDI types) and TWOSIDES (645 drugs, 200 DDI types) with three fixed "
         "random seeds each, following the same protocol as the original KnowDDI paper. "
         "For lightweight experiments, topology compression targets max_nodes_per_hop and max_links "
         "(60% retention), while parameter compression synchronously reduces emb_dim/gsl_rel_emb_dim/"
         "MLP_hidden_dim (50% retention), ensuring orthogonal and reproducible comparisons.}}"),
        ("{{Y:Specifically, on DrugBank the Denoise-only variant achieves the highest Macro-F1 (91.52%) "
         "among all four variants, Completion-only yields the highest ACC (92.89%) and Cohen's κ (91.57%), "
         "and the Denoise+Completion variant reaches peak ACC (92.99%) and κ (91.68%) but sees "
         "Macro-F1 drop to 90.78% due to long-tail class mis-completion edges. "
         "On TWOSIDES the baseline achieves AUROC 95.44% and AUPRC 94.11%, and all three GSL variants "
         "fall below it; the completion activation ratio self-converges from ~1.1% in early epochs "
         "to 0.6%–0.8% at convergence without manual intervention. "
         "For lightweight schemes on DrugBank, sparse-only causes 1.73% Macro-F1 loss with 43.9% "
         "memory savings, while co-compression offers no meaningful latency gain over feature-only "
         "because GraphSAGE embedding indexing latency scales linearly with emb_dim rather than "
         "node count; on TWOSIDES all three lightweight schemes yield less than 10% efficiency gain, "
         "indicating that compute cost is data-volume-driven rather than model-structure-driven.}}"),
    ]

    for base, ext in zip(texts_base, texts_ext):
        add_body_paragraph_with_highlights(doc, base + ext)  # 扩充英文摘要段落 -- 修改新增

    # 新增英文可解释性段落（绿色标注）
    text_en_new = (
        "Furthermore, two representative drug pairs—Reboxetine-Atomoxetine and Atropine-Scopolamine—"
        "are analyzed to demonstrate the interpretability of KnowDDI via knowledge subgraph learning. "
        "The former presents a clear two-hop reasoning path through a bridging compound (DB01146), "
        "while the latter reflects the model's multi-neighbor aggregation-based reasoning. "
        "These case studies validate that the knowledge subgraph serves not only as a structural "
        "tool for improving prediction accuracy, but also as a biomedical evidence carrier for "
        "interpreting model predictions."
    )
    p_en_new = doc.add_paragraph()
    p_en_new.paragraph_format.first_line_indent = Pt(24)
    p_en_new.paragraph_format.space_before = Pt(0)
    p_en_new.paragraph_format.space_after = Pt(0)
    r_en_new = p_en_new.add_run(text_en_new)
    set_font(r_en_new, "宋体", "Times New Roman", 12)
    r_en_new.font.color.rgb = GREEN  # 绿色标注新增内容

    doc.add_paragraph()
    kw_para = doc.add_paragraph()
    kw_para.paragraph_format.first_line_indent = Pt(0)
    r = kw_para.add_run("Keywords: ")
    set_font(r, "宋体", "Arial", 12, bold=True)
    r2 = kw_para.add_run("Drug-Drug Interaction Prediction; Graph Neural Network; Knowledge Graph; "
                          "Graph Structure Learning; Lightweight Model; Ablation Study; Interpretability")
    set_font(r2, "宋体", "Times New Roman", 12)

    add_page_break(doc)


# ──────────────────────────────────────────────
# 目录
# ──────────────────────────────────────────────

def add_toc(doc):
    """生成目录页（静态文本，建议在 Word 中刷新字段）"""
    add_chapter_heading(doc, "目  录")

    toc_entries = [
        ("摘要", ""),
        ("ABSTRACT", ""),
        ("第一章  绪论", ""),
        ("  1.1 研究背景与意义", ""),
        ("  1.2 国内外研究现状", ""),
        ("  1.3 研究内容与创新点", ""),
        ("  1.4 论文结构安排", ""),
        ("第二章  相关理论与技术基础", ""),
        ("  2.1 知识图谱与生物医学知识图谱", ""),
        ("  2.2 图神经网络基础", ""),
        ("  2.3 图结构学习", ""),
        ("  2.4 小样本学习背景下的DDI预测", ""),
        ("  2.5 模型轻量化技术", ""),
        ("第三章  KnowDDI模型分析与基准复现", ""),
        ("  3.1 KnowDDI模型整体架构", ""),
        ("  3.2 模型关键模块详解", ""),
        ("  3.3 实验环境与数据集准备", ""),
        ("  3.4 基准模型复现结果", ""),
        ("  3.5 典型药物对知识子图路径分析", ""),
        ("第四章  图结构学习机制的消融解耦研究", ""),
        ("  4.1 消融实验设计", ""),
        ("  4.2 DrugBank消融实验结果与分析", ""),
        ("  4.3 TWOSIDES消融实验结果与分析", ""),
        ("  4.4 跨数据集对比与综合分析", ""),
        ("第五章  面向边缘部署的轻量化重构与评测", ""),
        ("  5.1 轻量化方案设计", ""),
        ("  5.2 不同压缩比例下的性能预实验", ""),
        ("  5.3 正式轻量化方案评测结果", ""),
        ("  5.4 轻量化适用边界与折中关系总结", ""),
        ("  5.5 轻量化模型的可解释性验证", ""),
        ("第六章  总结与展望", ""),
        ("  6.1 工作总结", ""),
        ("  6.2 研究局限性", ""),
        ("  6.3 研究展望", ""),
        ("参考文献", ""),
        ("致谢", ""),
    ]

    for title, page in toc_entries:
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Pt(0)
        run = p.add_run(title)
        is_chapter = title.startswith("第") or title in ("摘要", "ABSTRACT", "参考文献", "致谢")
        set_font(run, "宋体", "Times New Roman", 12, bold=is_chapter)

    add_page_break(doc)


# ──────────────────────────────────────────────
# 参考文献
# ──────────────────────────────────────────────

def add_references(doc, refs_mod):
    add_chapter_heading(doc, "参考文献")
    for ref in refs_mod.REFERENCES:
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Pt(0)
        r = p.add_run(ref)
        set_font(r, "宋体", "Times New Roman", 10.5)
    add_page_break(doc)


# ──────────────────────────────────────────────
# 致谢
# ──────────────────────────────────────────────

def add_acknowledgement(doc, refs_mod):
    add_chapter_heading(doc, "致  谢")
    for line in refs_mod.ACKNOWLEDGEMENT.strip().split("\n"):
        if line.strip():
            add_body_paragraph_with_highlights(doc, line.strip())


# ──────────────────────────────────────────────
# 主构建函数
# ──────────────────────────────────────────────

def build_thesis():
    global _triggered_keywords
    _triggered_keywords = set()
    doc = Document()
    setup_page(doc)

    # 封面
    add_cover(doc)

    # 中英文摘要
    add_abstract_cn(doc)
    add_abstract_en(doc)

    # 目录
    add_toc(doc)

    # 各章节正文
    chapter_files = CHAPTER_FILES[:-1]  # 除了 references_acknowledgement
    for fname in chapter_files:
        mod = load_chapter_module(fname)
        # 获取新增节标记（若有）
        new_sections = getattr(mod, 'NEW_SECTIONS', set())
        # 章标题
        add_chapter_heading(doc, mod.CHAPTER_TITLE)

        for sec_num, sec_data in mod.SECTIONS.items():
            is_new = sec_num in new_sections  # 判断是否为新增节
            # 节标题：新增节用绿色
            sec_heading = add_section_heading(doc, f"{sec_num} {sec_data['title']}", level=2)
            if is_new:
                for run in sec_heading.runs:
                    run.font.color.rgb = GREEN  # 绿色标注新增节标题
            # 节正文
            render_section_content(doc, sec_data["content"], use_green=is_new)

        add_page_break(doc)

    # 参考文献 & 致谢
    refs_mod = load_chapter_module("references_acknowledgement.py")
    add_references(doc, refs_mod)
    add_acknowledgement(doc, refs_mod)

    doc.save(OUTPUT_PATH)
    print(f"[完成] 论文已生成：{OUTPUT_PATH}")


if __name__ == "__main__":
    build_thesis()
